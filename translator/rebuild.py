"""
Builds the output .docx from the original blocks + their translated text,
preserving paragraph order/structure, forcing Times New Roman throughout,
and setting right-to-left paragraph/run properties when the target language
is RTL (Urdu, Arabic, Farsi, Hebrew, etc.).
"""
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from . import prompts
from .translate import ChunkResult

FONT_NAME = "Times New Roman"


def _set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    for run in paragraph.runs:
        rPr = run._r.get_or_add_rPr()
        rtl = OxmlElement("w:rtl")
        rPr.append(rtl)


def _apply_font(run, rtl: bool):
    run.font.name = FONT_NAME
    # Required so Word uses the Latin font name for complex-script (RTL) runs too
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)  # complex script (Arabic/Urdu/Hebrew script)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)


def build_docx(
    chunk_results: list[ChunkResult],
    target_lang: str,
    output_path: str,
):
    rtl = prompts.is_rtl(target_lang)
    doc = Document()

    for chunk in chunk_results:
        for block, translated in zip(chunk.blocks, chunk.translated_texts):
            p = doc.add_paragraph()
            run = p.add_run(translated if translated else block["text"])
            run.font.size = Pt(12)
            if block.get("bold"):
                run.bold = True
            _apply_font(run, rtl)

            if block["style"] == "heading1":
                run.bold = True
                run.font.size = Pt(16)
            elif block["style"] == "heading2":
                run.bold = True
                run.font.size = Pt(13)
            elif block["style"] == "bullet":
                p.style = doc.styles["List Bullet"] if "List Bullet" in [s.name for s in doc.styles] else p.style

            if rtl:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _set_rtl(p)
            elif block.get("alignment"):
                # best-effort passthrough of original alignment for LTR docs
                pass

    doc.save(output_path)
    return output_path


def output_filename(original_filename: str, target_lang: str) -> str:
    base, _ext = os.path.splitext(os.path.basename(original_filename))
    safe_lang = target_lang.strip().replace(" ", "_")
    return f"{base}_{safe_lang}.docx"
