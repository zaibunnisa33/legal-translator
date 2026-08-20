"""
Extracts text as an ordered list of "blocks" from docx / pdf / jpeg (or png).
Each block is a dict:
    {
        "text": str,
        "style": "heading1" | "heading2" | "normal" | "bullet",
        "alignment": "left" | "center" | "right" | None,
        "bold": bool,
    }
Images (jpeg/png) and scanned/text-light PDFs are read via GPT-4o Vision OCR,
which is far more reliable on legal documents (stamps, tables, mixed fonts)
than classic OCR engines.
"""
import base64
import io
from docx import Document as DocxDocument
import pdfplumber
from PIL import Image

try:
    from openai import OpenAI  # only needed for OCR (pdf-scan/image) paths
except ImportError:
    OpenAI = None


def extract_docx(path: str) -> list[dict]:
    doc = DocxDocument(path)
    blocks = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = (para.style.name or "").lower()
        if "heading 1" in style_name or "title" in style_name:
            style = "heading1"
        elif "heading" in style_name:
            style = "heading2"
        elif "list" in style_name or para.text.strip().startswith(("-", "•")):
            style = "bullet"
        else:
            style = "normal"

        align = None
        if para.alignment is not None:
            align = str(para.alignment)

        bold = any(run.bold for run in para.runs if run.bold is not None)

        blocks.append({"text": text, "style": style, "alignment": align, "bold": bold})

    # Tables: flatten cell-by-cell, tagged so rebuild can at least preserve content
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    blocks.append({"text": t, "style": "normal", "alignment": None, "bold": False})

    return blocks


def _image_to_data_url(img: Image.Image) -> str:
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    b64 = base64.b64encode(buf.getvalue()).decode()
    return f"data:image/png;base64,{b64}"


def _ocr_image_with_gpt(client: OpenAI, img: Image.Image, model: str = "gpt-4o") -> str:
    data_url = _image_to_data_url(img)
    resp = client.chat.completions.create(
        model=model,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "Transcribe ALL text in this image exactly as written, preserving "
                            "line breaks, numbering, and paragraph structure. Do not translate, "
                            "summarise, or comment. Output only the transcribed text."
                        ),
                    },
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            }
        ],
        temperature=0,
    )
    return resp.choices[0].message.content.strip()


def extract_image(client: OpenAI, path: str, model: str = "gpt-4o") -> list[dict]:
    img = Image.open(path)
    text = _ocr_image_with_gpt(client, img, model)
    blocks = []
    for line in text.split("\n"):
        line = line.strip()
        if line:
            blocks.append({"text": line, "style": "normal", "alignment": None, "bold": False})
    return blocks


def extract_pdf(client: OpenAI, path: str, model: str = "gpt-4o") -> list[dict]:
    blocks = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                for line in text.split("\n"):
                    line = line.strip()
                    if line:
                        blocks.append({"text": line, "style": "normal", "alignment": None, "bold": False})
            else:
                # Likely a scanned page with no extractable text layer -> OCR it
                pil_img = page.to_image(resolution=200).original
                text = _ocr_image_with_gpt(client, pil_img, model)
                for line in text.split("\n"):
                    line = line.strip()
                    if line:
                        blocks.append({"text": line, "style": "normal", "alignment": None, "bold": False})
    return blocks


def chunk_blocks(blocks: list[dict], max_chars: int = 1800) -> list[list[dict]]:
    """Group consecutive blocks into chunks under max_chars, without splitting
    a block. Returns a list of lists-of-blocks (chunks)."""
    chunks = []
    current: list[dict] = []
    current_len = 0
    for b in blocks:
        blen = len(b["text"])
        if current and current_len + blen > max_chars:
            chunks.append(current)
            current = []
            current_len = 0
        current.append(b)
        current_len += blen
    if current:
        chunks.append(current)
    return chunks
